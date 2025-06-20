
# from langchain_google_genai import ChatGoogleGenerativeAI
# from langchain_core.prompts import ChatPromptTemplate
# from langchain_community.vectorstores import FAISS
# from langchain_community.document_loaders import PyPDFLoader
# from langchain_text_splitters import RecursiveCharacterTextSplitter
# from langchain_community.embeddings import SentenceTransformerEmbeddings
# from langgraph.graph import StateGraph, END
# from typing import TypedDict, Any, Optional, Dict, List
# import os
# from pathlib import Path
# import json
# import logging
# import asyncio
# import time
# from datetime import datetime
# from django.conf import settings
# from langchain_core.runnables import RunnableLambda
# import tempfile
# from docx import Document
# from langchain_core.output_parsers import JsonOutputParser
# from langchain.agents import create_react_agent, AgentExecutor as LangChainAgentExecutor
# from langchain_core.tools import tool
# from pydantic import BaseModel, Field
# from enum import Enum
# from dataclasses import dataclass

# # Set up logging with configurable level
# logging.basicConfig(level=os.getenv("LOG_LEVEL", "INFO"))
# logger = logging.getLogger(__name__)

# # Define execution status enum
# class ExecutionStatus(Enum):
#     PENDING = "pending"
#     RUNNING = "running"
#     COMPLETED = "completed"
#     FAILED = "failed"
#     CANCELLED = "cancelled"

# # Define execution result dataclass
# @dataclass
# class ExecutionResult:
#     status: ExecutionStatus
#     result: Optional[Dict[str, Any]] = None
#     error: Optional[str] = None
#     execution_time: Optional[float] = None
#     timestamp: Optional[datetime] = None
#     agent_logs: Optional[List[Dict]] = None

# # Define input validation schemas
# class AspirationsInput(BaseModel):
#     ctc: str = Field(description="Desired CTC, e.g., 12 LPA")
#     companies: str = Field(description="Target companies, e.g., Google")
#     role: str = Field(description="Preferred job role, e.g., software engineer")

# class ProfileInput(BaseModel):
#     coding: Optional[str] = Field(default="", description="Coding skills description")
#     experience: Optional[str] = Field(default="", description="Work experience details")
#     projects: Optional[str] = Field(default="", description="Project details")

# # Define state for the workflow
# class AgentState(TypedDict):
#     aspirations: dict
#     profile_score: dict
#     market_benchmarks: dict
#     reality_score: dict
#     micro_okrs: list
#     report: dict
#     resume_file: Any

# # ReAct tools
# @tool
# async def retrieve_market_data(aspirations: str) -> List[Dict]:
#     """Retrieve market data for given aspirations."""
#     retriever = AgentExecutor._get_retriever()
#     docs = await retriever.ainvoke(aspirations)
#     return [{"content": doc.page_content, "metadata": doc.metadata} for doc in docs]

# @tool
# async def evaluate_profile(profile: str, resume_text: str) -> Dict:
#     """Evaluate profile strength based on profile and resume data."""
#     prompt = ChatPromptTemplate.from_template("""
#     Analyze profile strength from: {profile} and resume: {resume_text}.
#     Score (0-100):
#     - Coding skills
#     - Work experience
#     - Projects
#     Output JSON with keys: coding, experience, projects.
#     """)
#     chain = prompt | AgentExecutor._get_llm() | JsonOutputParser()
#     return await chain.ainvoke({"profile": profile, "resume_text": resume_text})

# class AgentExecutor:
#     """
#     Enhanced Agent Executor with ReAct pattern and efficiency improvements
#     """
#     _retriever = None
#     _llm = None
#     _text_splitter = None

#     def __init__(self, max_retries: int = 3, timeout: int = 300):
#         self.max_retries = max_retries
#         self.timeout = timeout
#         self.execution_history: List[ExecutionResult] = []
#         self.current_execution: Optional[ExecutionResult] = None
#         self._setup_rag()
#         self._build_workflow()

#     @staticmethod
#     def _get_llm():
#         if AgentExecutor._llm is None:
#             google_api_key = getattr(settings, 'GOOGLE_API_KEY', None)
#             if not google_api_key or not google_api_key.strip():
#                 logger.error("GOOGLE_API_KEY not configured.")
#                 raise ValueError("GOOGLE_API_KEY not configured.")
#             AgentExecutor._llm = ChatGoogleGenerativeAI(
#                 model="gemini-1.5-flash",
#                 api_key=google_api_key,
#                 temperature=0.7
#             )
#         return AgentExecutor._llm

#     @staticmethod
#     def _get_retriever():
#         if AgentExecutor._retriever is None:
#             AgentExecutor._setup_rag()
#         return AgentExecutor._retriever

#     @staticmethod
#     def _setup_rag():
#         """Setup RAG components with dynamic paths and error handling"""
#         rag_dataset_path = os.getenv("RAG_DATASET_PATH", Path(settings.BASE_DIR) / "utils" / "RAG_Dataset.pdf")
#         if not rag_dataset_path.exists():
#             logger.error("RAG dataset not found at %s", rag_dataset_path)
#             raise FileNotFoundError(f"RAG dataset not found at {rag_dataset_path}")

#         try:
#             loader = PyPDFLoader(str(rag_dataset_path))
#             documents = loader.load()
#             AgentExecutor._text_splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=200)
#             texts = AgentExecutor._text_splitter.split_documents(documents)
#             embeddings = SentenceTransformerEmbeddings(model_name="all-MiniLM-L6-v2")
#             vectorstore = FAISS.from_documents(texts, embeddings)
#             AgentExecutor._retriever = vectorstore.as_retriever(search_kwargs={"k": 3})
#             logger.info("RAG setup completed.")
#         except Exception as e:
#             logger.error("Failed to set up RAG: %s", e)
#             raise

#     def _build_workflow(self):
#         """Build LangGraph workflow with ReAct agent"""
#         workflow = StateGraph(AgentState)
#         workflow.add_node("aspiration_parser", self._aspiration_parser_agent)
#         workflow.add_node("parallel_agents", self._parallel_agents)
#         workflow.add_node("compute_reality_score", self._reality_score_agent)
#         workflow.add_node("delta_reducer", self._delta_reducer_agent)
#         workflow.add_node("self_awareness_coach", self._self_awareness_coach_agent)

#         workflow.set_entry_point("aspiration_parser")
#         workflow.add_edge("aspiration_parser", "parallel_agents")
#         workflow.add_edge("parallel_agents", "compute_reality_score")
#         workflow.add_edge("compute_reality_score", "delta_reducer")
#         workflow.add_edge("delta_reducer", "self_awareness_coach")
#         workflow.add_edge("self_awareness_coach", END)

#         self.app = workflow.compile()

#     async def execute(
#         self,
#         aspirations: Dict[str, Any],
#         profile_data: Dict[str, Any],
#         resume_file: Optional[Any] = None,
#         async_execution: bool = True
#     ) -> ExecutionResult:
#         """Execute the agent workflow with ReAct and monitoring"""
#         try:
#             AspirationsInput(**aspirations)
#             ProfileInput(**profile_data)
#         except Exception as e:
#             logger.error("Invalid input: %s", e)
#             return ExecutionResult(
#                 status=ExecutionStatus.FAILED,
#                 error=f"Invalid input: {e}",
#                 timestamp=datetime.now()
#             )

#         if async_execution:
#             return await self._execute_async(aspirations, profile_data, resume_file)
#         else:
#             return self._execute_sync(aspirations, profile_data, resume_file)

#     def _execute_sync(
#         self,
#         aspirations: Dict[str, Any],
#         profile_data: Dict[str, Any],
#         resume_file: Optional[Any] = None
#     ) -> ExecutionResult:
#         """Synchronous execution"""
#         start_time = time.time()
#         execution_result = ExecutionResult(
#             status=ExecutionStatus.PENDING,
#             timestamp=datetime.now(),
#             agent_logs=None
#         )
#         self.current_execution = execution_result

#         try:
#             execution_result.status = ExecutionStatus.RUNNING
#             logger.info("Starting agent workflow execution")

#             initial_state = {
#                 "aspirations": aspirations,
#                 "profile": profile_data,
#                 "resume_file": resume_file
#             }
#             result = self._execute_with_retry(initial_state)

#             execution_result.result = {
#                 "aspirations": result.get("aspirations", {}),
#                 "profile_score": result.get("profile_score", {}),
#                 "market_benchmarks": result.get("market_benchmarks", {}),
#                 "gap_score": result.get("reality_score", {}).get("score", 0),
#                 "gap_details": result.get("reality_score", {}).get("details", {}),
#                 "micro_okrs": result.get("micro_okrs", []),
#                 "report": result.get("report", {})
#             }
#             execution_result.status = ExecutionStatus.COMPLETED
#             execution_result.execution_time = time.time() - start_time
#             logger.info(f"Workflow completed in {execution_result.execution_time:.2f} seconds")
#         except Exception as e:
#             execution_result.status = ExecutionStatus.FAILED
#             execution_result.error = str(e)
#             execution_result.execution_time = time.time() - start_time
#             execution_result.result = self._get_fallback_result(aspirations, profile_data)
#             logger.error(f"Workflow failed: {e}")
#         finally:
#             self.execution_history.append(execution_result)
#             self.current_execution = None
#         return execution_result

#     async def _execute_async(
#         self,
#         aspirations: Dict[str, Any],
#         profile_data: Dict[str, Any],
#         resume_file: Optional[Any] = None
#     ) -> ExecutionResult:
#         """Asynchronous execution with LangGraph"""
#         start_time = time.time()
#         execution_result = ExecutionResult(
#             status=ExecutionStatus.PENDING,
#             timestamp=datetime.now(),
#             agent_logs=None
#         )
#         self.current_execution = execution_result

#         try:
#             execution_result.status = ExecutionStatus.RUNNING
#             logger.info("Starting async agent workflow")

#             initial_state = {
#                 "aspirations": aspirations,
#                 "profile": profile_data,
#                 "resume_file": resume_file
#             }
#             result = await self._execute_with_retry_async(initial_state)

#             execution_result.result = {
#                 "aspirations": result.get("aspirations", {}),
#                 "profile_score": result.get("profile_score", {}),
#                 "market_benchmarks": result.get("market_benchmarks", {}),
#                 "gap_score": result.get("reality_score", {}).get("score", 0),
#                 "gap_details": result.get("reality_score", {}).get("details", {}),
#                 "micro_okrs": result.get("micro_okrs", []),
#                 "report": result.get("report", {})
#             }
#             execution_result.status = ExecutionStatus.COMPLETED
#             execution_result.execution_time = time.time() - start_time
#             logger.info(f"Async workflow completed in {execution_result.execution_time:.2f} seconds")
#         except Exception as e:
#             execution_result.status = ExecutionStatus.FAILED
#             execution_result.error = str(e)
#             execution_result.execution_time = time.time() - start_time
#             execution_result.result = self._get_fallback_result(aspirations, profile_data)
#             logger.error(f"Async workflow failed: {e}")
#         finally:
#             self.execution_history.append(execution_result)
#             self.current_execution = None
#         return execution_result

#     def _execute_with_retry(self, initial_state: Dict[str, Any]) -> Dict[str, Any]:
#         """Synchronous retry logic"""
#         last_exception = None
#         for attempt in range(self.max_retries):
#             try:
#                 logger.debug(f"Sync attempt {attempt + 1}/{self.max_retries}")
#                 return self.app.invoke(initial_state)
#             except Exception as e:
#                 last_exception = e
#                 if attempt < self.max_retries - 1:
#                     wait_time = 2 ** attempt
#                     logger.info(f"Retrying in {wait_time} seconds...")
#                     time.sleep(wait_time)
#         raise last_exception

#     async def _execute_with_retry_async(self, initial_state: Dict[str, Any]) -> Dict[str, Any]:
#         """Asynchronous retry logic"""
#         last_exception = None
#         for attempt in range(self.max_retries):
#             try:
#                 logger.debug(f"Async attempt {attempt + 1}/{self.max_retries}")
#                 return await self.app.ainvoke(initial_state)
#             except Exception as e:
#                 last_exception = e
#                 if attempt < self.max_retries - 1:
#                     wait_time = 2 ** attempt
#                     logger.info(f"Retrying in {wait_time} seconds...")
#                     await asyncio.sleep(wait_time)
#         raise last_exception

#     def _get_fallback_result(self, aspirations: Dict, profile_data: Dict) -> Dict[str, Any]:
#         """Dynamic fallback result based on input context"""
#         return {
#             "aspirations": aspirations,
#             "profile_score": {"coding": 0, "experience": 0, "projects": 0},
#             "market_benchmarks": {"ctc_range": "Not available", "skills": [], "experience": "Not available"},
#             "gap_score": 0,
#             "gap_details": {"ctc": "Analysis failed", "skills": "Analysis failed", "experience": "Analysis failed"},
#             "micro_okrs": [
#                 {"task": f"Learn core skills for {aspirations.get('role', 'target role')}", "resource": "Online courses", "timeline": "Week 1"},
#                 {"task": "Update professional profiles", "resource": "LinkedIn, resume guides", "timeline": "Week 2"}
#             ],
#             "report": {
#                 "summary": f"Unable to analyze for {aspirations.get('role', 'target role')}. Please verify inputs.",
#                 "feedback": "System error occurred. Check input data and retry.",
#                 "next_steps": "Review aspirations and profile data for accuracy."
#             }
#         }

#     def get_execution_status(self) -> Optional[ExecutionStatus]:
#         return self.current_execution.status if self.current_execution else None

#     def get_execution_history(self) -> List[ExecutionResult]:
#         return self.execution_history.copy()

#     def cancel_execution(self) -> bool:
#         if self.current_execution and self.current_execution.status == ExecutionStatus.RUNNING:
#             self.current_execution.status = ExecutionStatus.CANCELLED
#             logger.info("Execution cancelled")
#             return True
#         return False

#     async def process_resume_file(self, resume_file):
#         """Process resume file with enhanced error handling"""
#         if not resume_file:
#             return []
#         try:
#             with tempfile.NamedTemporaryFile(delete=False, suffix=os.path.splitext(resume_file.name)[1]) as temp_file:
#                 for chunk in resume_file.chunks():
#                     temp_file.write(chunk)
#                 temp_file_path = temp_file.name

#             if resume_file.name.endswith('.pdf'):
#                 loader = PyPDFLoader(temp_file_path)
#                 documents = await loader.aload()
#                 texts = AgentExecutor._text_splitter.split_documents(documents)
#             elif resume_file.name.endswith(('.doc', '.docx')):
#                 doc = Document(temp_file_path)
#                 text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
#                 texts = AgentExecutor._text_splitter.create_documents([text])
#             else:
#                 raise ValueError(f"Unsupported file format: {resume_file.name}")
#             return texts
#         except Exception as e:
#             logger.error("Error processing resume: %s", e)
#             raise
#         finally:
#             if 'temp_file_path' in locals():
#                 os.unlink(temp_file_path)

#     def _aspiration_parser_agent(self, state):
#         """ReAct-based aspiration parser"""
#         prompt = ChatPromptTemplate.from_template("""
#         You are a career coach analyzing user aspirations. Your task is to extract key details from the input provided.

#         Tools available: {tools}
#         Tool names: {tool_names}

#         Input aspirations: {input}

#         Instructions:
#         - Reason through the input to extract:
#           - Desired CTC (e.g., "12 LPA")
#           - Target companies (e.g., "Google, Amazon")
#           - Preferred job role (e.g., "software engineer")
#         - Use the retrieve_market_data tool if needed to infer reasonable defaults for ambiguous inputs.
#         - If the input is unclear, provide speculative defaults based on market trends.
#         - Output a valid JSON object with keys: ctc, companies, role.

#         Example output:
#         ```json
#         {{ "ctc": "12 LPA", "companies": "Google, Amazon", "role": "software engineer" }}
#         ```

#         Scratchpad for your reasoning:
#         {agent_scratchpad}

#         Respond only with the JSON output.
#         """)
#         react_agent = create_react_agent(
#             llm=self._get_llm(),
#             tools=[retrieve_market_data],
#             prompt=prompt
#         )
#         executor = LangChainAgentExecutor(agent=react_agent, tools=[retrieve_market_data], max_iterations=3)
#         try:
#             response = executor.invoke({"input": json.dumps(state.get("aspirations", {}))})
#             result = JsonOutputParser().parse(response["output"])
#             return {"aspirations": result}
#         except Exception as e:
#             logger.error("Aspiration parser error: %s", e)
#             return {"aspirations": state.get("aspirations", {"ctc": "Not extracted", "companies": "Not extracted", "role": "Not extracted"})}

#     async def _parallel_agents(self, state):
#         """Parallel execution of profile and market agents"""
#         tasks = [
#             self._profile_evaluator_agent(state),
#             self._market_benchmarking_agent(state)
#         ]
#         results = await asyncio.gather(*tasks, return_exceptions=True)
#         profile_score = (
#             results[0].get("profile_score", {"coding": 0, "experience": 0, "projects": 0})
#             if not isinstance(results[0], Exception)
#             else {"coding": 0, "experience": 0, "projects": 0}
#         )
#         market_benchmarks = (
#             results[1].get("market_benchmarks", {"ctc_range": "Not available", "skills": [], "experience": "Not available"})
#             if not isinstance(results[1], Exception)
#             else {"ctc_range": "Not available", "skills": [], "experience": "Not available"}
#         )
#         return {"profile_score": profile_score, "market_benchmarks": market_benchmarks}

#     async def _profile_evaluator_agent(self, state):
#         """Profile evaluator using LLM chain"""
#         resume_texts = await self.process_resume_file(state.get("resume_file")) if "resume_file" in state else []
#         profile_text = json.dumps(state.get("profile", {}))
#         resume_text = " ".join([doc.page_content for doc in resume_texts])
#         try:
#             prompt = ChatPromptTemplate.from_template("""
#             Analyze profile strength from: {profile} and resume: {resume_text}.
#             Score (0-100):
#             - Coding skills
#             - Work experience
#             - Projects
#             Output JSON with keys: coding, experience, projects.
#             """)
#             chain = prompt | self._get_llm() | JsonOutputParser()
#             result = await chain.ainvoke({"profile": profile_text, "resume_text": resume_text})
#             return {"profile_score": result}
#         except Exception as e:
#             logger.error("Profile evaluator error: %s", e)
#             return {"profile_score": {"coding": 0, "experience": 0, "projects": 0}}

#     async def _market_benchmarking_agent(self, state):
#         """Market benchmarking using retriever and LLM chain"""
#         try:
#             # Use retriever directly instead of tool
#             retriever = self._get_retriever()
#             aspirations_json = json.dumps(state.get("aspirations", {}))
#             docs = await retriever.ainvoke(aspirations_json)
#             doc_contents = [doc.page_content for doc in docs]
#             prompt = ChatPromptTemplate.from_template("""
#             Analyze market data: {docs} for aspirations: {aspirations}.
#             Output JSON: {{ "ctc_range": str, "skills": list, "experience": str }}
#             """)
#             chain = prompt | self._get_llm() | JsonOutputParser()
#             result = await chain.ainvoke({
#                 "aspirations": aspirations_json,
#                 "docs": doc_contents
#             })
#             return {"market_benchmarks": result}
#         except Exception as e:
#             logger.error("Market benchmarking error: %s", e)
#             return {"market_benchmarks": {"ctc_range": "Not available", "skills": [], "experience": "Not available"}}

#     async def _reality_score_agent(self, state):
#         """Compute reality score"""
#         prompt = ChatPromptTemplate.from_template("""
#         Score gap (0-100) between:
#         - Aspirations: {aspirations}
#         - Profile: {profile_score}
#         - Market: {market_benchmarks}
#         Output JSON: {{ "score": int, "details": {{ "ctc": str, "skills": str, "experience": str }} }}
#         """)
#         chain = prompt | self._get_llm() | JsonOutputParser()
#         try:
#             result = await chain.ainvoke({
#                 "aspirations": json.dumps(state.get("aspirations", {})),
#                 "profile_score": json.dumps(state.get("profile_score", {})),
#                 "market_benchmarks": json.dumps(state.get("market_benchmarks", {}))
#             })
#             return {"reality_score": result}
#         except Exception as e:
#             logger.error("Reality score error: %s", e)
#             return {"reality_score": {"score": 0, "details": {"ctc": "Not calculated", "skills": "Not calculated", "experience": "Not calculated"}}}

#     async def _delta_reducer_agent(self, state):
#         """Generate micro-OKRs"""
#         prompt = ChatPromptTemplate.from_template("""
#         Generate 30-day micro-OKR plan based on:
#         - Reality score: {reality_score}
#         - Profile: {profile_score}
#         - Market: {market_benchmarks}
#         - Aspirations: {aspirations}
#         Output JSON: {{ "okrs": list of {{ "task": str, "resource": str, "timeline": str }} }}
#         """)
#         chain = prompt | self._get_llm() | JsonOutputParser()
#         try:
#             result = await chain.ainvoke({
#                 "reality_score": json.dumps(state.get("reality_score", {})),
#                 "profile_score": json.dumps(state.get("profile_score", {})),
#                 "market_benchmarks": json.dumps(state.get("market_benchmarks", {})),
#                 "aspirations": json.dumps(state.get("aspirations", {}))
#             })
#             return {"micro_okrs": result.get("okrs", [])}
#         except Exception as e:
#             logger.error("Delta reducer error: %s", e)
#             return {"micro_okrs": [
#                 {"task": f"Learn skills for {state.get('aspirations', {}).get('role', 'target role')}", "resource": "Online courses", "timeline": "Week 1"},
#                 {"task": "Update profiles", "resource": "LinkedIn", "timeline": "Week 2"}
#             ]}

#     async def _self_awareness_coach_agent(self, state):
#         """Generate coaching report"""
#         prompt = ChatPromptTemplate.from_template("""
#         Generate coaching report based on:
#         - Aspirations: {aspirations}
#         - Profile: {profile_score}
#         - Market: {market_benchmarks}
#         - Reality score: {reality_score}
#         - Micro OKRs: {micro_okrs}
#         Output JSON: {{ "summary": str, "feedback": str, "next_steps": str }}
#         """)
#         chain = prompt | self._get_llm() | JsonOutputParser()
#         try:
#             result = await chain.ainvoke({
#                 "aspirations": json.dumps(state.get("aspirations", {})),
#                 "profile_score": json.dumps(state.get("profile_score", {})),
#                 "market_benchmarks": json.dumps(state.get("market_benchmarks", {})),
#                 "reality_score": json.dumps(state.get("reality_score", {})),
#                 "micro_okrs": json.dumps(state.get("micro_okrs", []))
#             })
#             return {"report": result}
#         except Exception as e:
#             logger.error("Coach error: %s", e)
#             return {"report": self._generate_personalized_fallback(state)}

#     def _generate_personalized_fallback(self, state):
#         """Dynamic fallback report"""
#         aspirations = state.get("aspirations", {})
#         return {
#             "summary": f"Unable to analyze for {aspirations.get('role', 'target role')}.",
#             "feedback": "Check input data and retry.",
#             "next_steps": "Verify aspirations and profile data."
#         }

# # Global executor instance
# executor = AgentExecutor()

# # Enhanced execution function
# def execute_agent_workflow(aspirations, profile_data, resume_file=None, async_execution=True):
#     return executor.execute(aspirations, profile_data, resume_file, async_execution)

# # Legacy function for backward compatibility
# def run_agent_workflow(aspirations, profile_data, resume_file=None):
#     """Legacy function for backward compatibility"""
#     result = asyncio.run(executor.execute(aspirations, profile_data, resume_file, async_execution=True))
#     if result.status == ExecutionStatus.COMPLETED:
#         return result.result
#     else:
#         return {
#             "error": result.error or "Failed to process request.",
#             "details": result.error
#         }

# def get_execution_status():
#     return executor.get_execution_status()

# def get_execution_history():
#     return executor.get_execution_history()

# def cancel_execution():
#     return executor.cancel_execution()

from langchain_google_genai import ChatGoogleGenerativeAI
from langchain_core.prompts import ChatPromptTemplate
from langchain_community.vectorstores import FAISS
from langchain_community.document_loaders import PyPDFLoader
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.embeddings import SentenceTransformerEmbeddings
from typing import Any, Optional, Dict, List
import os
from pathlib import Path
import json
import logging
import time
from datetime import datetime
from django.conf import settings
from langchain_core.output_parsers import JsonOutputParser
from langchain.agents import create_react_agent, AgentExecutor
from langchain_core.tools import tool
from pydantic import BaseModel, Field
from enum import Enum
from dataclasses import dataclass
import tempfile
from docx import Document
import threading
import queue

# Set up logging with configurable level
logging.basicConfig(level=os.getenv("LOG_LEVEL", "INFO"))
logger = logging.getLogger(__name__)

# Define execution status enum
class ExecutionStatus(Enum):
    PENDING = "pending"
    RUNNING = "running"
    COMPLETED = "completed"
    FAILED = "failed"
    CANCELLED = "cancelled"

# Define execution result dataclass
@dataclass
class ExecutionResult:
    status: ExecutionStatus
    result: Optional[Dict[str, Any]] = None
    error: Optional[str] = None
    execution_time: Optional[float] = None
    timestamp: Optional[datetime] = None
    agent_logs: Optional[List[Dict]] = None

# Define input validation schemas
class AspirationsInput(BaseModel):
    ctc: str = Field(description="Desired CTC, e.g., 12 LPA")
    companies: str = Field(description="Target companies, e.g., Google")
    role: str = Field(description="Preferred job role, e.g., software engineer")

class ProfileInput(BaseModel):
    coding: Optional[str] = Field(default="", description="Coding skills description")
    experience: Optional[str] = Field(default="", description="Work experience details")
    projects: Optional[str] = Field(default="", description="Project details")

# Initialize LLM and retriever
class AgentComponents:
    _llm = None
    _retriever = None
    _text_splitter = None

    @staticmethod
    def get_llm():
        if AgentComponents._llm is None:
            google_api_key = getattr(settings, 'GOOGLE_API_KEY', None)
            if not google_api_key or not google_api_key.strip():
                logger.error("GOOGLE_API_KEY not configured.")
                raise ValueError("GOOGLE_API_KEY not configured.")
            AgentComponents._llm = ChatGoogleGenerativeAI(
                model="gemini-1.5-flash",
                api_key=google_api_key,
                temperature=0.7
            )
        return AgentComponents._llm

    @staticmethod
    def get_retriever():
        if AgentComponents._retriever is None:
            AgentComponents.setup_rag()
        return AgentComponents._retriever

    @staticmethod
    def get_text_splitter():
        if AgentComponents._text_splitter is None:
            AgentComponents.setup_rag()
        return AgentComponents._text_splitter

    @staticmethod
    def setup_rag():
        """Setup RAG components with dynamic paths and error handling"""
        rag_dataset_path = os.getenv("RAG_DATASET_PATH", str(Path(settings.BASE_DIR) / "utils" / "RAG_Dataset.pdf"))
        rag_dataset_path = Path(rag_dataset_path)  # Ensure it's a Path object
        if not rag_dataset_path.exists():
            logger.error("RAG dataset not found at %s", rag_dataset_path)
            raise FileNotFoundError(f"RAG dataset not found at {rag_dataset_path}")

        try:
            loader = PyPDFLoader(str(rag_dataset_path))
            documents = loader.load()
            AgentComponents._text_splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=200)
            texts = AgentComponents._text_splitter.split_documents(documents)
            embeddings = SentenceTransformerEmbeddings(model_name="all-MiniLM-L6-v2")
            vectorstore = FAISS.from_documents(texts, embeddings)
            AgentComponents._retriever = vectorstore.as_retriever(search_kwargs={"k": 3})
            logger.info("RAG setup completed.")
        except Exception as e:
            logger.error("Failed to set up RAG: %s", e)
            raise

# ReAct tools
@tool
def retrieve_market_data(aspirations: str) -> List[Dict]:
    """Retrieve market data for given aspirations."""
    retriever = AgentComponents.get_retriever()
    docs = retriever.invoke(aspirations)
    return [{"content": doc.page_content, "metadata": doc.metadata} for doc in docs]

@tool
def evaluate_profile(profile: str, resume_text: str) -> Dict:
    """Evaluate profile strength based on profile and resume data."""
    prompt = ChatPromptTemplate.from_template("""
    Analyze profile strength from: {profile} and resume: {resume_text}.
    Score (0-100):
    - Coding skills
    - Work experience
    - Projects
    Output JSON with keys: coding, experience, projects.
    """)
    chain = prompt | AgentComponents.get_llm() | JsonOutputParser()
    return chain.invoke({"profile": profile, "resume_text": resume_text})

def process_resume_file(resume_file):
    """Process resume file with enhanced error handling"""
    if not resume_file:
        return []
    try:
        with tempfile.NamedTemporaryFile(delete=False, suffix=os.path.splitext(resume_file.name)[1]) as temp_file:
            for chunk in resume_file.chunks():
                temp_file.write(chunk)
            temp_file_path = temp_file.name

        if resume_file.name.endswith('.pdf'):
            loader = PyPDFLoader(temp_file_path)
            documents = loader.load()
            text_splitter = AgentComponents.get_text_splitter()
            texts = text_splitter.split_documents(documents)
        elif resume_file.name.endswith(('.doc', '.docx')):
            doc = Document(temp_file_path)
            text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
            text_splitter = AgentComponents.get_text_splitter()
            texts = text_splitter.create_documents([text])
        else:
            raise ValueError(f"Unsupported file format: {resume_file.name}")
        return texts
    except Exception as e:
        logger.error("Error processing resume: %s", e)
        raise
    finally:
        if 'temp_file_path' in locals():
            os.unlink(temp_file_path)

def execute_agent_workflow(aspirations: Dict[str, Any], profile_data: Dict[str, Any], resume_file: Optional[Any] = None) -> ExecutionResult:
    """Execute the agent workflow using LangChain AgentExecutor"""
    start_time = time.time()
    execution_result = ExecutionResult(
        status=ExecutionStatus.PENDING,
        timestamp=datetime.now(),
        agent_logs=[]
    )

    try:
        # Validate inputs
        AspirationsInput(**aspirations)
        ProfileInput(**profile_data)
    except Exception as e:
        logger.error("Invalid input: %s", e)
        return ExecutionResult(
            status=ExecutionStatus.FAILED,
            error=f"Invalid input: {e}",
            timestamp=datetime.now(),
            agent_logs=execution_result.agent_logs
        )

    execution_result.status = ExecutionStatus.RUNNING
    logger.info("Starting agent workflow execution")

    try:
        # Step 1: Parse aspirations using ReAct agent
        prompt = ChatPromptTemplate.from_template("""
        You are a career coach analyzing user aspirations. Your task is to extract key details from the input provided.

        Tools available: {tools}
        Tool names: {tool_names}

        Input aspirations: {input}

        Instructions:
        - Follow the ReAct (Reasoning-Action) format: provide 'Thought', 'Action', and 'Final Answer' steps.
        - Reason through the input to extract:
          - Desired CTC (e.g., "12 LPA")
          - Target companies (e.g., "Google, Amazon")
          - Preferred job role (e.g., "software engineer")
        - If the input is clear, extract the details directly.
        - If the input is unclear or incomplete, use the retrieve_market_data tool to infer reasonable defaults based on market trends.
        - Provide step-by-step reasoning in the scratchpad under 'Thought'.
        - Use 'Action' to either extract details or call the retrieve_market_data tool.
        - Output the final result in the 'Final Answer' section as a valid JSON object with keys: ctc, companies, role.
        - Ensure the JSON output is properly formatted and enclosed in a 'Final Answer' block.

        Example output:
        ```
        Thought: The input provides clear details about CTC, companies, and role. No need to use the tool.
        Action: Extract details from input.
        Final Answer: {{"ctc": "12 LPA", "companies": "Google, Amazon", "role": "software engineer"}}
        ```

        Scratchpad for your reasoning:
        {agent_scratchpad}
        """)
        react_agent = create_react_agent(
            llm=AgentComponents.get_llm(),
            tools=[retrieve_market_data],
            prompt=prompt
        )
        aspiration_executor = AgentExecutor(
            agent=react_agent,
            tools=[retrieve_market_data],
            max_iterations=5,
            verbose=True,
            handle_parsing_errors=True,
            return_intermediate_steps=True
        )
        response = aspiration_executor.invoke({"input": json.dumps(aspirations)})
        parsed_aspirations = JsonOutputParser().parse(response["output"])
        execution_result.agent_logs.append({
            "agent": "aspiration_parser",
            "input": aspirations,
            "output": parsed_aspirations,
            "intermediate_steps": response.get("intermediate_steps", [])
        })

        # Step 2: Parallel execution of profile and market analysis
        profile_queue = queue.Queue()
        market_queue = queue.Queue()

        def profile_evaluator():
            try:
                resume_texts = process_resume_file(resume_file) if resume_file else []
                profile_text = json.dumps(profile_data)
                resume_text = " ".join([doc.page_content for doc in resume_texts])
                prompt = ChatPromptTemplate.from_template("""
                Analyze profile strength from: {profile} and resume: {resume_text}.
                Score (0-100):
                - Coding skills
                - Work experience
                - Projects
                Output JSON with keys: coding, experience, projects.
                """)
                chain = prompt | AgentComponents.get_llm() | JsonOutputParser()
                result = chain.invoke({"profile": profile_text, "resume_text": resume_text})
                execution_result.agent_logs.append({
                    "agent": "profile_evaluator",
                    "input": {"profile": profile_text, "resume_text": resume_text},
                    "output": result
                })
                profile_queue.put(result)
            except Exception as e:
                logger.error("Profile evaluator error: %s", e)
                execution_result.agent_logs.append({
                    "agent": "profile_evaluator",
                    "error": str(e)
                })
                profile_queue.put({"coding": 0, "experience": 0, "projects": 0})

        def market_benchmarking():
            try:
                retriever = AgentComponents.get_retriever()
                aspirations_json = json.dumps(parsed_aspirations)
                docs = retriever.invoke(aspirations_json)
                doc_contents = [doc.page_content for doc in docs]
                prompt = ChatPromptTemplate.from_template("""
                Analyze market data: {docs} for aspirations: {aspirations}.
                Output JSON: {{ "ctc_range": str, "skills": list, "experience": str }}
                """)
                chain = prompt | AgentComponents.get_llm() | JsonOutputParser()
                result = chain.invoke({"aspirations": aspirations_json, "docs": doc_contents})
                execution_result.agent_logs.append({
                    "agent": "market_benchmarking",
                    "input": {"aspirations": aspirations_json, "docs": doc_contents},
                    "output": result
                })
                market_queue.put(result)
            except Exception as e:
                logger.error("Market benchmarking error: %s", e)
                execution_result.agent_logs.append({
                    "agent": "market_benchmarking",
                    "error": str(e)
                })
                market_queue.put({"ctc_range": "Not available", "skills": [], "experience": "Not available"})

        profile_thread = threading.Thread(target=profile_evaluator)
        market_thread = threading.Thread(target=market_benchmarking)
        profile_thread.start()
        market_thread.start()
        profile_thread.join()
        market_thread.join()
        profile_score = profile_queue.get()
        market_benchmarks = market_queue.get()

        # Step 3: Compute reality score
        prompt = ChatPromptTemplate.from_template("""
        Score gap (0-100) between:
        - Aspirations: {aspirations}
        - Profile: {profile_score}
        - Market: {market_benchmarks}
        Output JSON: {{ "score": int, "details": {{ "ctc": str, "skills": str, "experience": str }} }}
        """)
        chain = prompt | AgentComponents.get_llm() | JsonOutputParser()
        try:
            reality_score = chain.invoke({
                "aspirations": json.dumps(parsed_aspirations),
                "profile_score": json.dumps(profile_score),
                "market_benchmarks": json.dumps(market_benchmarks)
            })
            execution_result.agent_logs.append({
                "agent": "reality_score",
                "input": {
                    "aspirations": parsed_aspirations,
                    "profile_score": profile_score,
                    "market_benchmarks": market_benchmarks
                },
                "output": reality_score
            })
        except Exception as e:
            logger.error("Reality score error: %s", e)
            reality_score = {
                "score": 0,
                "details": {"ctc": "Not calculated", "skills": "Not calculated", "experience": "Not calculated"}
            }
            execution_result.agent_logs.append({
                "agent": "reality_score",
                "error": str(e)
            })

        # Step 4: Generate micro-OKRs
        prompt = ChatPromptTemplate.from_template("""
        Generate 30-day micro-OKR plan based on:
        - Reality score: {reality_score}
        - Profile: {profile_score}
        - Market: {market_benchmarks}
        - Aspirations: {aspirations}
        Output JSON: {{ "okrs": list of {{ "task": str, "resource": str, "timeline": str }} }}
        Ensure at least five OKRs are provided, tailored to address specific gaps.
        """)
        chain = prompt | AgentComponents.get_llm() | JsonOutputParser()
        try:
            micro_okrs_result = chain.invoke({
                "reality_score": json.dumps(reality_score),
                "profile_score": json.dumps(profile_score),
                "market_benchmarks": json.dumps(market_benchmarks),
                "aspirations": json.dumps(parsed_aspirations)
            })
            micro_okrs = micro_okrs_result.get("okrs", [])
            if len(micro_okrs) < 5:
                logger.warning("LLM generated fewer than 5 OKRs, supplementing with defaults")
                micro_okrs.extend([
                    {"task": f"Learn core skills for {parsed_aspirations.get('role', 'target role')}", "resource": "Online courses (e.g., Coursera, Udemy)", "timeline": "Week 1"},
                    {"task": f"Complete 20 coding problems for {parsed_aspirations.get('role', 'target role')}", "resource": "LeetCode, HackerRank", "timeline": "Week 1-2"},
                    {"task": "Build a portfolio project relevant to role", "resource": "GitHub, project tutorials", "timeline": "Week 2-3"},
                    {"task": "Update professional profiles and resume", "resource": "LinkedIn, resume guides", "timeline": "Week 3"},
                    {"task": "Network with professionals in target companies", "resource": "LinkedIn, industry events", "timeline": "Week 4"}
                ][:5 - len(micro_okrs)])
            execution_result.agent_logs.append({
                "agent": "delta_reducer",
                "input": {
                    "reality_score": reality_score,
                    "profile_score": profile_score,
                    "market_benchmarks": market_benchmarks,
                    "aspirations": parsed_aspirations
                },
                "output": micro_okrs
            })
        except Exception as e:
            logger.error("Delta reducer error: %s", e)
            micro_okrs = [
                {"task": f"Learn core skills for {parsed_aspirations.get('role', 'target role')}", "resource": "Online courses (e.g., Coursera, Udemy)", "timeline": "Week 1"},
                {"task": f"Complete 20 coding problems for {parsed_aspirations.get('role', 'target role')}", "resource": "LeetCode, HackerRank", "timeline": "Week 1-2"},
                {"task": "Build a portfolio project relevant to role", "resource": "GitHub, project tutorials", "timeline": "Week 2-3"},
                {"task": "Update professional profiles and resume", "resource": "LinkedIn, resume guides", "timeline": "Week 3"},
                {"task": "Network with professionals in target companies", "resource": "LinkedIn, industry events", "timeline": "Week 4"}
            ]
            execution_result.agent_logs.append({
                "agent": "delta_reducer",
                "error": str(e)
            })

        # Step 5: Generate coaching report
        prompt = ChatPromptTemplate.from_template("""
        Generate coaching report based on:
        - Aspirations: {aspirations}
        - Profile: {profile_score}
        - Market: {market_benchmarks}
        - Reality score: {reality_score}
        - Micro OKRs: {micro_okrs}
        Output JSON: {{ "summary": str, "feedback": str, "next_steps": str }}
        Use a conversational, empathetic tone.
        """)
        chain = prompt | AgentComponents.get_llm() | JsonOutputParser()
        try:
            report = chain.invoke({
                "aspirations": json.dumps(parsed_aspirations),
                "profile_score": json.dumps(profile_score),
                "market_benchmarks": json.dumps(market_benchmarks),
                "reality_score": json.dumps(reality_score),
                "micro_okrs": json.dumps(micro_okrs)
            })
            execution_result.agent_logs.append({
                "agent": "self_awareness_coach",
                "input": {
                    "aspirations": parsed_aspirations,
                    "profile_score": profile_score,
                    "market_benchmarks": market_benchmarks,
                    "reality_score": reality_score,
                    "micro_okrs": micro_okrs
                },
                "output": report
            })
        except Exception as e:
            logger.error("Coach error: %s", e)
            report = {
                "summary": f"Unable to analyze for {parsed_aspirations.get('role', 'target role')}.",
                "feedback": "Check input data and retry.",
                "next_steps": "Verify aspirations and profile data."
            }
            execution_result.agent_logs.append({
                "agent": "self_awareness_coach",
                "error": str(e)
            })

        # Compile final result
        execution_result.result = {
            "aspirations": parsed_aspirations,
            "profile_score": profile_score,
            "market_benchmarks": market_benchmarks,
            "gap_score": reality_score.get("score", 0),
            "gap_details": reality_score.get("details", {}),
            "micro_okrs": micro_okrs,
            "report": report
        }
        execution_result.status = ExecutionStatus.COMPLETED
        execution_result.execution_time = time.time() - start_time
        logger.info(f"Workflow completed in {execution_result.execution_time:.2f} seconds")
    except Exception as e:
        execution_result.status = ExecutionStatus.FAILED
        execution_result.error = str(e)
        execution_result.execution_time = time.time() - start_time
        execution_result.result = {
            "aspirations": aspirations,
            "profile_score": {"coding": 0, "experience": 0, "projects": 0},
            "market_benchmarks": {"ctc_range": "Not available", "skills": [], "experience": "Not available"},
            "gap_score": 0,
            "gap_details": {"ctc": "Analysis failed", "skills": "Analysis failed", "experience": "Analysis failed"},
            "micro_okrs": [
                {"task": f"Learn core skills for {aspirations.get('role', 'target role')}", "resource": "Online courses (e.g., Coursera, Udemy)", "timeline": "Week 1"},
                {"task": f"Complete 20 coding problems for {aspirations.get('role', 'target role')}", "resource": "LeetCode, HackerRank", "timeline": "Week 1-2"},
                {"task": "Build a portfolio project relevant to role", "resource": "GitHub, project tutorials", "timeline": "Week 2-3"},
                {"task": "Update professional profiles and resume", "resource": "LinkedIn, resume guides", "timeline": "Week 3"},
                {"task": "Network with professionals in target companies", "resource": "LinkedIn, industry events", "timeline": "Week 4"}
            ],
            "report": {
                "summary": f"Unable to analyze for {aspirations.get('role', 'target role')}. Please verify inputs.",
                "feedback": "System error occurred. Check input data and retry.",
                "next_steps": "Review aspirations and profile data for accuracy."
            }
        }
        logger.error(f"Workflow failed: {e}")
    
    return execution_result

# Legacy function for backward compatibility
def run_agent_workflow(aspirations, profile_data, resume_file=None):
    """Legacy function for backward compatibility"""
    result = execute_agent_workflow(aspirations, profile_data, resume_file)
    if result.status == ExecutionStatus.COMPLETED:
        return result.result
    else:
        return {
            "error": result.error or "Failed to process request.",
            "details": result.error
        }